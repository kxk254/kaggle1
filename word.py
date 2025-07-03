from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Word文書の作成
doc = Document()

# スタイル設定
style = doc.styles['Normal']
font = style.font
font.name = 'MS Gothic'
font.size = Pt(10.5)

# 日付・宛名
p1 = doc.add_paragraph('　　　　　　　　　　　　　　　　　　　　　　　　　　　　　令和　　年　　月　　日')
p1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

doc.add_paragraph('合同会社 CS Tsukiji 御中')
doc.add_paragraph('（ご担当者様）')

doc.add_paragraph('\n　　　　　　　　　　　　　　　　　　　　　　　　【貴社名】')
doc.add_paragraph('　　　　　　　　　　　　　　　　　　　　　　　　【ご担当者名】 様')

doc.add_paragraph('\n下記の通り、工事作業のご報告を申し上げます。\n')

# 工事概要テーブル
doc.add_heading('■ 工事概要', level=2)

table_data = [
    ['項目', '内容'],
    ['工事名', '（　　　　　　　　　　　　　）'],
    ['管理ビル名', '築地センタービル'],
    ['工事場所', '築地センタービル ５階'],
    ['実施日', '（　　　年　　月　　日）'],
    ['作業時間', '（　　時　　分 ～ 　　時　　分）'],
    ['作業責任者', '（　　　　　　　　　　　　　）'],
    ['ご連絡先（電話番号）', '（　　　　　　　　　　　　　）'],
    ['作業人数', '（　　名）'],
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '項目'
hdr_cells[1].text = '内容'

for row in table_data[1:]:
    cells = table.add_row().cells
    cells[0].text = row[0]
    cells[1].text = row[1]

# 注意事項
doc.add_paragraph('\n')
doc.add_heading('■ ご注意事項（遵守のお願い）', level=2)

doc.add_paragraph(
    '以下の内容につき、作業時に必ずご対応いただきますようお願い申し上げます。\n\n'
    '・作業エリアの養生（養生マット、ブルーシート等）を実施すること\n'
    '・共用部の通行制限は基本行わず、必要な場合は警備員を配置すること\n'
    '・火気使用は原則禁止（使用時は事前申請が必要）\n'
    '・騒音作業は 午前9時～午後5時まで の時間帯に限定すること\n'
    '・作業終了後は、清掃・原状回復を徹底すること\n'
    '・エレベーターの使用は、12時～13時の昼休み時間を避けること\n'
    '・工程表、作業図面 を添付のうえ、提出願います'
)

doc.add_paragraph('\nご確認のほどよろしくお願い申し上げます。')

# 保存
doc.save('工事作業書_築地センタービル.docx')
print("Wordファイル '工事作業書_築地センタービル.docx' を作成しました。")
